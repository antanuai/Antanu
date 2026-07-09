# -*- coding: utf-8 -*-
"""
export_utils.py — ساخت خروجی Word و PDF فارسی (راست‌به‌چپ) با فونت و سایز دلخواه
"""
import os
import re
import secrets

# پوشه خروجی‌ها — اگر قابل نوشتن نبود (مثل Hugging Face) به /tmp می‌رود
_BASE = os.path.dirname(os.path.abspath(__file__))
EXPORT_DIR = os.environ.get("ANTANU_EXPORT_DIR", os.path.join(_BASE, "exports"))
try:
    os.makedirs(EXPORT_DIR, exist_ok=True)
    _t = os.path.join(EXPORT_DIR, ".w")
    open(_t, "w").close()
    os.remove(_t)
except Exception:
    EXPORT_DIR = "/tmp/antanu_exports"
    os.makedirs(EXPORT_DIR, exist_ok=True)

FONT_DIR = os.path.join(_BASE, "static", "fonts")
os.makedirs(FONT_DIR, exist_ok=True)
PDF_FONT_PATH = os.path.join(FONT_DIR, "Vazirmatn-Regular.ttf")
FONT_URLS = [
    "https://cdn.jsdelivr.net/gh/rastikerdar/vazirmatn@v33.003/fonts/ttf/Vazirmatn-Regular.ttf",
    "https://raw.githubusercontent.com/rastikerdar/vazirmatn/v33.003/fonts/ttf/Vazirmatn-Regular.ttf",
]


def ensure_pdf_font() -> str | None:
    """فونت فارسی برای PDF — اگر نبود، بار اول دانلود می‌شود"""
    if os.path.exists(PDF_FONT_PATH) and os.path.getsize(PDF_FONT_PATH) > 50_000:
        return PDF_FONT_PATH
    try:
        import httpx
        for url in FONT_URLS:
            try:
                r = httpx.get(url, timeout=30, follow_redirects=True)
                if r.status_code == 200 and len(r.content) > 50_000:
                    with open(PDF_FONT_PATH, "wb") as f:
                        f.write(r.content)
                    return PDF_FONT_PATH
            except Exception:
                continue
    except Exception:
        pass
    return None


# ---------------- تبدیل مارک‌داون ساده به بلوک‌ها ----------------

def md_to_blocks(text: str):
    """('h1'|'h2'|'h3'|'li'|'p', متن) — علامت‌های مارک‌داون حذف می‌شوند"""
    blocks = []
    for raw in (text or "").split("\n"):
        line = raw.strip()
        if not line or line in ("---", "***"):
            continue
        line = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        line = re.sub(r"[*_`]", "", line)
        if line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
        elif line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
        elif line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
        elif re.match(r"^[-•]\s+", line):
            blocks.append(("li", re.sub(r"^[-•]\s+", "", line)))
        elif re.match(r"^\d+[.)]\s+", line):
            blocks.append(("li", re.sub(r"^\d+[.)]\s+", "", line)))
        else:
            blocks.append(("p", line))
    return blocks


def _new_name(ext: str) -> str:
    return f"antanu-{secrets.token_hex(6)}.{ext}"


# ---------------- ساخت فایل Word (راست‌به‌چپ) ----------------

def build_docx(blocks, font_name: str = "Vazirmatn", font_size: int = 14, title: str | None = None) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def style_run(run, size, bold=False, color=None):
        run.font.name = font_name
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rFonts.set(qn(attr), font_name)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), str(int(size * 2)))
        rPr.append(szCs)
        rPr.append(OxmlElement("w:rtl"))

    def rtl_para(p, align=WD_ALIGN_PARAGRAPH.RIGHT):
        p.alignment = align
        pPr = p._p.get_or_add_pPr()
        pPr.append(OxmlElement("w:bidi"))

    doc = Document()

    if title:
        p = doc.add_paragraph()
        rtl_para(p, WD_ALIGN_PARAGRAPH.CENTER)
        style_run(p.add_run(title), font_size + 10, bold=True, color=(0x0F, 0x76, 0x6E))

    for kind, txt in blocks:
        p = doc.add_paragraph()
        rtl_para(p)
        if kind == "h1":
            style_run(p.add_run(txt), font_size + 8, bold=True, color=(0x0F, 0x76, 0x6E))
        elif kind == "h2":
            style_run(p.add_run(txt), font_size + 4, bold=True, color=(0xB8, 0x86, 0x0B))
        elif kind == "h3":
            style_run(p.add_run(txt), font_size + 2, bold=True)
        elif kind == "li":
            style_run(p.add_run("• " + txt), font_size)
        else:
            style_run(p.add_run(txt), font_size)

    name = _new_name("docx")
    doc.save(os.path.join(EXPORT_DIR, name))
    return name


# ---------------- ساخت فایل PDF فارسی ----------------

def build_pdf(blocks, font_size: int = 14, title: str | None = None):
    """خروجی: (نام فایل، None) یا (None، پیام خطا)"""
    try:
        from fpdf import FPDF
        import arabic_reshaper
        from bidi.algorithm import get_display
    except ImportError:
        return None, "برای PDF این کتابخانه‌ها لازم است: pip install fpdf2 arabic-reshaper python-bidi"

    font = ensure_pdf_font()
    if not font:
        return None, "فونت فارسی PDF دانلود نشد (اینترنت سرور را بررسی کنید) — فایل Word ساخته شد"

    def shape(t):
        try:
            return get_display(arabic_reshaper.reshape(t))
        except Exception:
            return t

    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()
    pdf.add_font("Vazir", "", font)

    if title:
        pdf.set_font("Vazir", size=font_size + 8)
        pdf.multi_cell(0, (font_size + 8) * 0.62, shape(title), align="C")
        pdf.ln(3)

    for kind, txt in blocks:
        size = font_size + (8 if kind == "h1" else 4 if kind == "h2" else 2 if kind == "h3" else 0)
        pdf.set_font("Vazir", size=size)
        text = ("• " + txt) if kind == "li" else txt
        try:
            pdf.multi_cell(0, size * 0.62, shape(text), align="R")
        except Exception:
            continue
        if kind in ("h1", "h2", "h3"):
            pdf.ln(1)

    name = _new_name("pdf")
    pdf.output(os.path.join(EXPORT_DIR, name))
    return name, None
