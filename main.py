# -*- coding: utf-8 -*-
"""
main.py — هسته اصلی هوش مصنوعی آنتانو (ANTANU)
اجرا:  uvicorn main:app --host 0.0.0.0 --port 8000
"""
import os
import re
import json
import asyncio
import secrets

import httpx
from fastapi import FastAPI, Request, Form, HTTPException, UploadFile, File
from starlette.concurrency import run_in_threadpool
from fastapi.responses import HTMLResponse, RedirectResponse, StreamingResponse, JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from jinja2 import Environment, FileSystemLoader, select_autoescape

from db import get_db, init_db, hash_pw, verify_pw, generate_code

# ---------------- بارگذاری فایل .env ----------------
# کلید API و تنظیمات محرمانه در فایل .env نگهداری می‌شوند (کنار main.py).
# این فایل هرگز نباید روی گیت‌هاب آپلود شود — در .gitignore مستثنی شده است.

def _load_env_file(path: str = ".env"):
    if not os.path.exists(path):
        return
    with open(path, encoding="utf-8-sig") as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, _, value = line.partition("=")
            key = key.strip()
            value = value.strip().strip('"').strip("'")
            if key and key not in os.environ:
                os.environ[key] = value


_load_env_file()

# ---------------- تنظیمات ----------------
# آنتانو با هر سرویس سازگار با OpenAI کار می‌کند: Groq، GitHub Models، Gemini، OpenRouter و...
# فقط کافی است ANTANU_PROVIDER و ANTANU_API_KEY را تنظیم کنید (راهنما در README).

PROVIDERS = {
    # سریع و رایگان — پیشنهادی برای تعداد کاربر زیاد
    "groq": {
        "base": "https://api.groq.com/openai/v1",
        "model": "llama-3.3-70b-versatile",
    },
    # رایگان با توکن گیت‌هاب (محدودیت روزانه کم دارد)
    "github": {
        "base": "https://models.github.ai/inference",
        "model": "openai/gpt-4o-mini",
    },
    # بهترین کیفیت فارسی در بین گزینه‌های رایگان
    "gemini": {
        "base": "https://generativelanguage.googleapis.com/v1beta/openai",
        "model": "gemini-2.0-flash",
    },
    # چت جی‌پی‌تی رسمی (پولی)
    "openai": {
        "base": "https://api.openai.com/v1",
        "model": "gpt-4o-mini",
    },
    # چندین مدل رایگان — openrouter/free خودش مدل رایگانِ در دسترس را انتخاب می‌کند
    "openrouter": {
        "base": "https://openrouter.ai/api/v1",
        "model": "openrouter/free, meta-llama/llama-3.3-70b:free, openai/gpt-oss-120b:free",
    },
}

PROVIDER = os.environ.get("ANTANU_PROVIDER", "groq").lower()
_p = PROVIDERS.get(PROVIDER, PROVIDERS["groq"])
API_BASE = os.environ.get("ANTANU_API_BASE", _p["base"]).rstrip("/")
API_KEY = os.environ.get("ANTANU_API_KEY", "")
# می‌توان چند مدل را با کاما جدا کرد؛ اگر اولی در دسترس نبود خودکار سراغ بعدی می‌رود
MODELS = [m.strip() for m in os.environ.get("ANTANU_MODEL", _p["model"]).split(",") if m.strip()]
ADMIN_CONTACT = os.environ.get("ANTANU_ADMIN_CONTACT", "آیدی تلگرام ادمین: @your_admin_id")

# ---------------- فهرست هوش مصنوعی‌های قابل انتخاب توسط کاربر ----------------
# دو روش تعریف:
# روش ساده: فقط ANTANU_PROVIDER و ANTANU_API_KEY را بدهید → فهرست پیش‌فرض همان سرویس ساخته می‌شود.
# روش چندسرویسه: در فایل .env هر هوش مصنوعی را در یک خط تعریف کنید (تا ۱۰ عدد):
#   ANTANU_AI_1=نام نمایشی | سرویس | نام مدل | کلید
#   «سرویس» یکی از این‌هاست: groq / gemini / github / openrouter / یا آدرس کامل API
# مثال:
#   ANTANU_AI_1=جمنای گوگل | gemini | gemini-2.0-flash | AIza...
#   ANTANU_AI_2=لاما (گراک) | groq | llama-3.3-70b-versatile | gsk_...
#   ANTANU_AI_3=دیپ‌سیک | openrouter | deepseek/deepseek-r1-distill:free | sk-or-...

def _parse_custom_ais():
    out = []
    for i in range(1, 11):
        raw = os.environ.get(f"ANTANU_AI_{i}", "").strip()
        if not raw:
            continue
        parts = [p.strip() for p in raw.split("|")]
        if len(parts) != 4:
            print(f"[ANTANU] قالب ANTANU_AI_{i} نادرست است — باید ۴ بخش جداشده با | باشد")
            continue
        name, prov, model, key = parts
        base = PROVIDERS[prov.lower()]["base"] if prov.lower() in PROVIDERS else prov.rstrip("/")
        out.append({"id": f"ai{i}", "name": name, "base": base, "model": model, "key": key})
    return out


_customs = _parse_custom_ais()

AI_CATALOG = []
if API_KEY:
    # «آنتانو (خودکار)» — هوش مصنوعی اصلی خود سایت
    AI_CATALOG.append(
        {"id": "auto", "name": "آنتانو (خودکار)", "base": API_BASE, "model": MODELS[0], "key": API_KEY}
    )
    if PROVIDER == "openrouter" and not _customs:
        # فهرست پیش‌فرض مدل‌های رایگان OpenRouter (با همان یک کلید)
        AI_CATALOG += [
            {"id": "llama",    "name": "Llama 3.3 70B",          "base": API_BASE, "model": "meta-llama/llama-3.3-70b:free",    "key": API_KEY},
            {"id": "gptoss",   "name": "GPT-OSS 120B (OpenAI)",  "base": API_BASE, "model": "openai/gpt-oss-120b:free",         "key": API_KEY},
            {"id": "deepseek", "name": "DeepSeek R1",            "base": API_BASE, "model": "deepseek/deepseek-r1-distill:free","key": API_KEY},
            {"id": "gptnano",  "name": "GPT-5.4 Nano (سریع)",    "base": API_BASE, "model": "openai/gpt-5.4-nano:free",         "key": API_KEY},
        ]
    elif not _customs and len(MODELS) > 1:
        AI_CATALOG += [
            {"id": f"m{i}", "name": m.split("/")[-1], "base": API_BASE, "model": m, "key": API_KEY}
            for i, m in enumerate(MODELS[1:], 1)
        ]

# هوش مصنوعی‌های تعریف‌شده توسط مدیر — دقیقاً همین‌ها به کاربر نمایش داده می‌شوند
AI_CATALOG += _customs

if not AI_CATALOG:
    AI_CATALOG = [{"id": "auto", "name": "آنتانو", "base": API_BASE, "model": MODELS[0], "key": ""}]

# کاتالوگ ساخته‌شده از فایل .env (به‌عنوان پشتیبان)
ENV_CATALOG = AI_CATALOG


def resolve_base(service: str) -> str:
    """نام سرویس (groq/gemini/...) یا آدرس کامل → آدرس پایه API"""
    s = (service or "").strip()
    if s.lower() in PROVIDERS:
        return PROVIDERS[s.lower()]["base"]
    return s.rstrip("/")


def get_ai_catalog():
    """فهرست نهایی هوش مصنوعی‌ها — اولویت با تنظیماتی است که ادمین در پنل ذخیره کرده"""
    from db import get_db as _gdb
    try:
        db = _gdb()
        row = db.execute("SELECT value FROM settings WHERE key = 'ai_config'").fetchone()
        db.close()
    except Exception:
        row = None
    if row:
        try:
            cfg = json.loads(row["value"])
        except json.JSONDecodeError:
            cfg = None
        if cfg:
            catalog = []
            main_key = (cfg.get("api_key") or "").strip()
            if main_key:
                prov = (cfg.get("provider") or "groq").lower()
                model = (cfg.get("model") or "").strip() or PROVIDERS.get(prov, PROVIDERS["groq"])["model"].split(",")[0].strip()
                catalog.append({
                    "id": "auto", "name": "آنتانو (خودکار)",
                    "base": resolve_base(prov), "model": model, "key": main_key,
                })
            for i, ai in enumerate(cfg.get("ais") or [], 1):
                if not (ai.get("key") or "").strip() or not (ai.get("model") or "").strip():
                    continue
                catalog.append({
                    "id": f"ai{i}",
                    "name": (ai.get("name") or f"مدل {i}").strip(),
                    "base": resolve_base(ai.get("service")),
                    "model": ai["model"].strip(),
                    "key": ai["key"].strip(),
                })
            if catalog:
                # اگر کلید اصلی خالی بود، «آنتانو (خودکار)» با اولین هوش مصنوعیِ دارای کلید کار کند
                if catalog[0]["id"] != "auto":
                    first = catalog[0]
                    catalog.insert(0, {"id": "auto", "name": "آنتانو (خودکار)",
                                       "base": first["base"], "model": first["model"], "key": first["key"]})
                return catalog
    # پشتیبان .env — اگر آن هم بی‌کلید بود ولی مدل دیگری کلید داشت، از آن استفاده کن
    cat = [dict(c) for c in ENV_CATALOG]
    if cat and not cat[0].get("key"):
        withkey = next((c for c in cat if c.get("key")), None)
        if withkey:
            cat[0].update({"base": withkey["base"], "model": withkey["model"], "key": withkey["key"]})
    return cat

# سقف پیام روزانه بر اساس ستاره اشتراک (None یعنی بدون محدودیت)
DAILY_LIMITS = {1: 30, 2: 80, 3: 200, 4: None}
# حداکثر طول پاسخ مدل بر اساس ستاره (-۱ یعنی آزاد)
MAX_TOKENS = {1: 450, 2: 900, 3: 1600, 4: -1}

app = FastAPI(title="ANTANU")
app.mount("/static", StaticFiles(directory="static"), name="static")

# رندر مستقیم قالب‌ها با Jinja2 (مستقل از نسخه starlette — بدون خطای ناسازگاری)
_jinja = Environment(
    loader=FileSystemLoader("templates"),
    autoescape=select_autoescape(["html"]),
)


def render(name: str, status_code: int = 200, **context) -> HTMLResponse:
    html = _jinja.get_template(name).render(**context)
    return HTMLResponse(html, status_code=status_code)


init_db()

BASE_SYSTEM_PROMPT = (
    "تو «آنتانو» (ANTANU) هستی؛ دستیار هوشمند فارسی‌زبان برای دانشجویان و پژوهشگران. "
    "قواعد نگارش که همیشه باید رعایت کنی: "
    "۱) به فارسیِ معیار، روان و طبیعی بنویس؛ از ترجمه تحت‌اللفظی و جمله‌بندی انگلیسی‌مآب جداً پرهیز کن. "
    "۲) دستور زبان، املا و نشانه‌گذاری فارسی را کامل رعایت کن: نیم‌فاصله در «می‌شود» و «کتاب‌ها»، فعل در انتهای جمله، حروف اضافه درست. "
    "۳) اصطلاحات تخصصی را به فارسی بنویس و در اولین اشاره، معادل انگلیسی را داخل پرانتز بیاور. "
    "۴) در حوزه‌های دانشگاهی — به‌ویژه حسابداری، مدیریت، آمار و روش تحقیق — دقیق و علمی پاسخ بده. "
    "۵) اگر چیزی را نمی‌دانی صادقانه بگو نمی‌دانم و حدس نزن. "
    "۶) پاسخ را ساختارمند ارائه کن و اگر از نتایج جستجوی وب استفاده کردی، منبع را ذکر کن. "
    "۷) از تکرار واژه‌ها، عبارت‌ها و مطالب پرهیز کن؛ همیشه واژگان متنوع و مطالب تازه به کار ببر."
)


# ---------------- ابزارهای احراز هویت ----------------

def current_user(request: Request):
    token = request.cookies.get("antanu_session")
    if not token:
        return None
    db = get_db()
    row = db.execute(
        "SELECT u.* FROM sessions s JOIN users u ON u.id = s.user_id WHERE s.token = ?",
        (token,),
    ).fetchone()
    db.close()
    return row


def require_user(request: Request):
    user = current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="ابتدا وارد حساب خود شوید")
    return user


def require_admin(request: Request):
    user = require_user(request)
    if not user["is_admin"]:
        raise HTTPException(status_code=403, detail="دسترسی فقط برای مدیر")
    return user


def make_session(user_id: int) -> str:
    token = secrets.token_hex(32)
    db = get_db()
    db.execute("INSERT INTO sessions (token, user_id) VALUES (?, ?)", (token, user_id))
    db.commit()
    db.close()
    return token


# ---------------- صفحات ----------------

@app.get("/", response_class=HTMLResponse)
def root(request: Request):
    return RedirectResponse("/chat" if current_user(request) else "/login")


@app.get("/login", response_class=HTMLResponse)
def login_page(request: Request):
    if current_user(request):
        return RedirectResponse("/chat")
    return render("login.html", error=None)


@app.post("/login", response_class=HTMLResponse)
def login(request: Request, username: str = Form(...), password: str = Form(...), fingerprint: str = Form("")):
    db = get_db()
    user = db.execute("SELECT * FROM users WHERE username = ?", (username.strip(),)).fetchone()

    def fail(msg):
        db.close()
        return render("login.html", error=msg, status_code=400)

    if not user or not verify_pw(password, user["password"]):
        return fail("نام کاربری یا گذرواژه نادرست است.")

    # قفل یک‌دستگاهی: هر حساب فقط روی همان دستگاهی که اولین‌بار وارد شده کار می‌کند (ادمین معاف است)
    if not user["is_admin"]:
        if user["device_fp"] and fingerprint and user["device_fp"] != fingerprint:
            return fail("این حساب به دستگاه دیگری متصل است. برای انتقال به دستگاه جدید با ادمین تماس بگیرید.")
        if not user["device_fp"] and fingerprint:
            db.execute("UPDATE users SET device_fp = ? WHERE id = ?", (fingerprint, user["id"]))
            db.commit()

    db.close()
    token = make_session(user["id"])
    resp = RedirectResponse("/chat", status_code=303)
    resp.set_cookie("antanu_session", token, httponly=True, samesite="lax", max_age=60 * 60 * 24 * 30)
    return resp


@app.get("/register", response_class=HTMLResponse)
def register_page(request: Request):
    if current_user(request):
        return RedirectResponse("/chat")
    return render("register.html", error=None)


@app.post("/register", response_class=HTMLResponse)
def register(
    request: Request,
    username: str = Form(...),
    password: str = Form(...),
    code: str = Form(...),
    fingerprint: str = Form(""),
):
    username, code = username.strip(), code.strip()
    db = get_db()

    def fail(msg):
        db.close()
        return render("register.html", error=msg, status_code=400)

    if len(username) < 3:
        return fail("نام کاربری باید حداقل ۳ حرف باشد.")
    if len(password) < 6:
        return fail("گذرواژه باید حداقل ۶ حرف باشد.")
    if not fingerprint:
        return fail("شناسه دستگاه دریافت نشد. صفحه را نوسازی کنید.")
    if db.execute("SELECT 1 FROM users WHERE username = ?", (username,)).fetchone():
        return fail("این نام کاربری قبلاً گرفته شده است.")

    code_row = db.execute("SELECT * FROM codes WHERE code = ?", (code,)).fetchone()
    if not code_row:
        return fail("کد ثبت‌نام نامعتبر است.")
    if code_row["used"]:
        return fail("این کد قبلاً استفاده شده است. هر کد فقط برای یک کاربر و یک دستگاه معتبر است.")

    cur = db.execute(
        "INSERT INTO users (username, password, stars, device_fp, code_used) VALUES (?, ?, ?, ?, ?)",
        (username, hash_pw(password), code_row["stars"], fingerprint, code),
    )
    user_id = cur.lastrowid
    db.execute("UPDATE codes SET used = 1, used_by = ? WHERE id = ?", (username, code_row["id"]))
    db.commit()
    db.close()

    token = make_session(user_id)
    resp = RedirectResponse("/chat", status_code=303)
    resp.set_cookie("antanu_session", token, httponly=True, samesite="lax", max_age=60 * 60 * 24 * 30)
    return resp


@app.post("/logout")
def logout(request: Request):
    token = request.cookies.get("antanu_session")
    if token:
        db = get_db()
        db.execute("DELETE FROM sessions WHERE token = ?", (token,))
        db.commit()
        db.close()
    resp = RedirectResponse("/login", status_code=303)
    resp.delete_cookie("antanu_session")
    return resp


@app.get("/buy", response_class=HTMLResponse)
def buy_page(request: Request):
    return render("buy.html", contact=ADMIN_CONTACT)


@app.get("/chat", response_class=HTMLResponse)
def chat_page(request: Request):
    user = current_user(request)
    if not user:
        return RedirectResponse("/login")
    limit = DAILY_LIMITS.get(user["stars"])
    return render("chat.html", user=user, daily_limit=limit if limit else "نامحدود")


@app.get("/admin", response_class=HTMLResponse)
def admin_page(request: Request):
    user = current_user(request)
    if not user:
        return RedirectResponse("/login")
    if not user["is_admin"]:
        return RedirectResponse("/chat")
    return render("admin.html", user=user)


# ---------------- API گفتگوها ----------------

@app.get("/api/conversations")
def list_conversations(request: Request):
    user = require_user(request)
    db = get_db()
    rows = db.execute(
        "SELECT id, title, created_at FROM conversations WHERE user_id = ? ORDER BY id DESC",
        (user["id"],),
    ).fetchall()
    db.close()
    return [dict(r) for r in rows]


@app.get("/api/conversations/{conv_id}/messages")
def conversation_messages(conv_id: int, request: Request):
    user = require_user(request)
    db = get_db()
    conv = db.execute(
        "SELECT id FROM conversations WHERE id = ? AND user_id = ?", (conv_id, user["id"])
    ).fetchone()
    if not conv:
        db.close()
        raise HTTPException(404, "گفتگو یافت نشد")
    rows = db.execute(
        "SELECT role, content FROM messages WHERE conversation_id = ? ORDER BY id", (conv_id,)
    ).fetchall()
    db.close()
    return [dict(r) for r in rows]


@app.delete("/api/conversations/{conv_id}")
def delete_conversation(conv_id: int, request: Request):
    user = require_user(request)
    db = get_db()
    db.execute("DELETE FROM conversations WHERE id = ? AND user_id = ?", (conv_id, user["id"]))
    db.commit()
    db.close()
    return {"ok": True}


# ---------------- API حافظه بلندمدت ----------------

@app.get("/api/memories")
def list_memories(request: Request):
    user = require_user(request)
    db = get_db()
    rows = db.execute(
        "SELECT id, content, created_at FROM memories WHERE user_id = ? ORDER BY id DESC",
        (user["id"],),
    ).fetchall()
    db.close()
    return [dict(r) for r in rows]


@app.post("/api/memory")
async def save_memory(request: Request):
    user = require_user(request)
    body = await request.json()
    content = (body.get("content") or "").strip()
    if not content:
        raise HTTPException(400, "متن خالی است")
    db = get_db()
    db.execute("INSERT INTO memories (user_id, content) VALUES (?, ?)", (user["id"], content[:4000]))
    db.commit()
    db.close()
    return {"ok": True}


@app.delete("/api/memories/{mem_id}")
def delete_memory(mem_id: int, request: Request):
    user = require_user(request)
    db = get_db()
    db.execute("DELETE FROM memories WHERE id = ? AND user_id = ?", (mem_id, user["id"]))
    db.commit()
    db.close()
    return {"ok": True}


# ---------------- جستجوی وب (رایگان و بدون کلید) ----------------

def _search_web_sync(query: str) -> str:
    """جستجوی وب با موتور DuckDuckGo — رایگان، بدون نیاز به کلید"""
    from ddgs import DDGS
    lines = []
    with DDGS() as d:
        for r in d.text(query, max_results=5):
            title = r.get("title", "")
            body = r.get("body", "")
            href = r.get("href", "")
            lines.append(f"- {title}: {body}\n  منبع: {href}")
    return "\n".join(lines)


# ---------------- API فهرست مدل‌ها و آپلود فایل ----------------

@app.get("/api/models")
def list_models(request: Request):
    require_user(request)
    return [{"id": c["id"], "name": c["name"]} for c in get_ai_catalog()]


ALLOWED_TEXT_EXT = {".txt", ".md", ".csv", ".json", ".py", ".html", ".xml", ".log"}


@app.post("/api/upload")
async def upload_file(request: Request, file: UploadFile = File(...)):
    user = require_user(request)
    raw = await file.read()
    if len(raw) > 10 * 1024 * 1024:
        raise HTTPException(400, "حجم فایل نباید بیشتر از ۱۰ مگابایت باشد")

    name = file.filename or "file"
    ext = os.path.splitext(name)[1].lower()
    text = ""

    try:
        if ext in ALLOWED_TEXT_EXT:
            text = raw.decode("utf-8", errors="ignore")
        elif ext == ".pdf":
            import io
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(raw))
            text = "\n".join((p.extract_text() or "") for p in reader.pages)
        elif ext == ".docx":
            import io
            from docx import Document
            doc = Document(io.BytesIO(raw))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(c.text for c in row.cells))
            text = "\n".join(parts)
        else:
            raise HTTPException(
                400,
                "این نوع فایل فعلاً پشتیبانی نمی‌شود. فرمت‌های مجاز: txt، pdf، docx، csv، md، json",
            )
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(400, "خواندن محتوای فایل ممکن نشد. از سالم بودن فایل مطمئن شوید.")

    text = text.strip()
    if not text:
        raise HTTPException(400, "متنی در این فایل پیدا نشد.")

    db = get_db()
    cur = db.execute(
        "INSERT INTO uploads (user_id, filename, content) VALUES (?, ?, ?)",
        (user["id"], name[:200], text[:60000]),
    )
    db.commit()
    upload_id = cur.lastrowid
    db.close()
    return {"id": upload_id, "filename": name, "chars": len(text)}


# ---------------- API چت (استریم — چندمدلی + جستجوی وب + فایل) ----------------

def build_system_prompt(user, memories) -> str:
    prompt = BASE_SYSTEM_PROMPT
    if user["stars"] < 4:
        prompt += " پاسخ‌ها را نسبتاً خلاصه و مفید ارائه بده."
    else:
        prompt += " پاسخ‌ها را کامل، عمیق و جامع ارائه بده."
    if memories:
        prompt += "\n\nحافظه بلندمدت این کاربر (همیشه در نظر بگیر):\n"
        prompt += "\n".join(f"- {m['content']}" for m in memories)
    return prompt


class ModelError(Exception):
    def __init__(self, status: int, body: str = ""):
        self.status = status
        self.body = body


async def stream_model(messages, stars: int, model: str, base: str, key: str):
    """استریم پاسخ از هر سرویس سازگار با OpenAI — هر مدل با آدرس و کلید خودش"""
    payload = {"model": model, "messages": messages, "stream": True}
    if MAX_TOKENS.get(stars, -1) > 0:
        payload["max_tokens"] = MAX_TOKENS[stars]

    headers = {"Authorization": f"Bearer {key}", "Content-Type": "application/json"}

    async with httpx.AsyncClient(timeout=httpx.Timeout(120, connect=15)) as client:
        async with client.stream(
            "POST", f"{base}/chat/completions", json=payload, headers=headers
        ) as r:
            if r.status_code != 200:
                body = (await r.aread()).decode(errors="ignore")[:400]
                raise ModelError(r.status_code, body)
            async for line in r.aiter_lines():
                if not line.startswith("data:"):
                    continue
                data = line[5:].strip()
                if data == "[DONE]":
                    break
                try:
                    obj = json.loads(data)
                except json.JSONDecodeError:
                    continue
                choices = obj.get("choices") or [{}]
                chunk = (choices[0].get("delta") or {}).get("content") or ""
                if chunk:
                    yield chunk


@app.post("/api/chat")
async def api_chat(request: Request):
    user = require_user(request)
    body = await request.json()
    message = (body.get("message") or "").strip()
    conv_id = body.get("conversation_id")
    selected_ids = body.get("models") or ["auto"]
    web_on = bool(body.get("web"))
    research = bool(body.get("research"))
    attachment_ids = body.get("attachments") or []
    if not message:
        raise HTTPException(400, "پیام خالی است")

    # ---------- دستور ذخیره در حافظه:  \save متن  یا  /save ----------
    save_match = re.match(r"^[\\/]\s*save\b[:\s]*", message, re.IGNORECASE)
    if save_match:
        remainder = message[save_match.end():].strip()
        db = get_db()
        if conv_id:
            conv = db.execute(
                "SELECT id FROM conversations WHERE id = ? AND user_id = ?", (conv_id, user["id"])
            ).fetchone()
            if not conv:
                db.close()
                raise HTTPException(404, "گفتگو یافت نشد")
        else:
            cur = db.execute(
                "INSERT INTO conversations (user_id, title) VALUES (?, ?)",
                (user["id"], "🧠 ذخیره در حافظه"),
            )
            conv_id = cur.lastrowid
        db.execute(
            "INSERT INTO messages (conversation_id, role, content) VALUES (?, 'user', ?)",
            (conv_id, message),
        )
        if not remainder:
            row = db.execute(
                "SELECT content FROM messages WHERE conversation_id = ? AND role = 'assistant' "
                "ORDER BY id DESC LIMIT 1",
                (conv_id,),
            ).fetchone()
            remainder = (row["content"] if row else "").strip()
        if remainder:
            db.execute(
                "INSERT INTO memories (user_id, content) VALUES (?, ?)",
                (user["id"], remainder[:4000]),
            )
            reply = "🧠 در حافظه بلندمدت آنتانو ذخیره شد و هرگز فراموش نمی‌شود."
        else:
            reply = "⚠️ چیزی برای ذخیره پیدا نشد. بنویسید: \\save متن موردنظر — یا بعد از پاسخ ربات فقط \\save بفرستید تا همان پاسخ ذخیره شود."
        db.execute(
            "INSERT INTO messages (conversation_id, role, content) VALUES (?, 'assistant', ?)",
            (conv_id, reply),
        )
        db.commit()
        db.close()

        async def save_gen():
            yield reply

        return StreamingResponse(
            save_gen(),
            media_type="text/plain; charset=utf-8",
            headers={"X-Conversation-Id": str(conv_id)},
        )

    # مدل‌های انتخاب‌شده توسط کاربر (یکی یا چندتا)
    catalog = get_ai_catalog()
    chosen = [c for c in catalog if c["id"] in selected_ids]
    if not chosen:
        chosen = [catalog[0]]

    # حالت تحقیق گروهی: همه هوش مصنوعی‌ها دست‌به‌دست هم پژوهش را کامل می‌کنند + جستجوی وب
    if research:
        chosen = list(catalog)
        web_on = True

    db = get_db()

    # بررسی سقف پیام روزانه بر اساس اشتراک
    limit = DAILY_LIMITS.get(user["stars"])
    if limit is not None:
        used = db.execute(
            """SELECT COUNT(*) AS c FROM messages m
               JOIN conversations c2 ON c2.id = m.conversation_id
               WHERE c2.user_id = ? AND m.role = 'user' AND date(m.created_at) = date('now')""",
            (user["id"],),
        ).fetchone()["c"]
        if used >= limit:
            db.close()
            raise HTTPException(
                429,
                f"سقف {limit} پیام روزانه اشتراک {user['stars']} ستاره شما تمام شد. "
                "برای ادامه، اشتراک بالاتر تهیه کنید یا فردا برگردید.",
            )

    # گفتگو
    if conv_id:
        conv = db.execute(
            "SELECT id FROM conversations WHERE id = ? AND user_id = ?", (conv_id, user["id"])
        ).fetchone()
        if not conv:
            db.close()
            raise HTTPException(404, "گفتگو یافت نشد")
    else:
        cur = db.execute(
            "INSERT INTO conversations (user_id, title) VALUES (?, ?)",
            (user["id"], message[:60]),
        )
        conv_id = cur.lastrowid

    # محتوای فایل‌های پیوست
    extra_context = ""
    attach_names = []
    for aid in attachment_ids[:5]:
        row = db.execute(
            "SELECT filename, content FROM uploads WHERE id = ? AND user_id = ?",
            (int(aid), user["id"]),
        ).fetchone()
        if row:
            attach_names.append(row["filename"])
            extra_context += f"\n\n[محتوای فایل پیوست «{row['filename']}»]:\n{row['content'][:20000]}"

    # جستجوی وب
    web_note = ""
    if web_on:
        try:
            results = await run_in_threadpool(_search_web_sync, message[:300])
            if results:
                extra_context += (
                    "\n\n[نتایج جستجوی وب — برای پاسخ به‌روز از این‌ها استفاده کن و منبع را ذکر کن]:\n"
                    + results
                )
        except Exception:
            web_note = "\n\n_🌐 جستجوی وب در دسترس نبود؛ پاسخ از دانش خود مدل است._"

    # پیام نمایشی که در تاریخچه ذخیره می‌شود
    display = message
    if attach_names:
        display += "\n📎 " + "، ".join(attach_names)
    if web_on:
        display += "\n🌐 با جستجوی وب"

    db.execute(
        "INSERT INTO messages (conversation_id, role, content) VALUES (?, 'user', ?)",
        (conv_id, display),
    )
    db.commit()

    history = db.execute(
        "SELECT role, content FROM messages WHERE conversation_id = ? ORDER BY id DESC LIMIT 24",
        (conv_id,),
    ).fetchall()[::-1]
    memories = db.execute(
        "SELECT content FROM memories WHERE user_id = ? ORDER BY id DESC LIMIT 30", (user["id"],)
    ).fetchall()
    db.close()

    msgs = [{"role": "system", "content": build_system_prompt(user, memories)}]
    msgs += [{"role": r["role"], "content": r["content"]} for r in history[:-1]]
    msgs.append({"role": "user", "content": message + extra_context})

    stars = user["stars"]

    async def gen():
        full = ""
        try:
            multi = len(chosen) > 1
            for idx, c in enumerate(chosen):
                if await request.is_disconnected():
                    break
                if multi:
                    tag = " — ادامه پژوهش" if (research and idx > 0) else ""
                    hdr = ("\n\n---\n\n" if full else "") + f"### 🤖 {c['name']}{tag}\n\n"
                    full += hdr
                    yield hdr

                # در تحقیق گروهی، مدل‌های بعدی متن قبلی را می‌گیرند و فقط مطالب تازه اضافه می‌کنند
                cur_msgs = msgs
                if research and idx > 0 and full:
                    cur_msgs = [
                        {"role": "system", "content": build_system_prompt(user, memories)
                         + " تو ادامه‌دهنده یک پژوهش گروهی هستی."},
                        {"role": "user", "content":
                            f"موضوع پژوهش: {message}\n\n"
                            f"متن نوشته‌شده تاکنون توسط پژوهشگران قبلی:\n{full[-6000:]}\n\n"
                            "ادامه پژوهش را بنویس: فقط بخش‌ها و مطالب کاملاً تازه با عنوان‌های جدید اضافه کن. "
                            "از تکرار واژه‌ها، جمله‌ها و مطالب قبلی جداً پرهیز کن و واژگان نو به کار ببر."},
                    ]

                if not c.get("key"):
                    note = f"⚠️ کلید «{c['name']}» تنظیم نشده است. مدیر سیستم باید در فایل .env کلید بگذارد."
                    full += note
                    yield note
                    continue

                got_any = False
                try:
                    async for chunk in stream_model(cur_msgs, stars, c["model"], c["base"], c["key"]):
                        got_any = True
                        full += chunk
                        yield chunk
                except ModelError as e:
                    # تلاش پشتیبان با مدل خودکار (فقط تک‌مدلی و فقط OpenRouter)
                    recovered = False
                    if (not multi and not got_any and "openrouter.ai" in c["base"]
                            and c["model"] != "openrouter/free"):
                        try:
                            async for chunk in stream_model(cur_msgs, stars, "openrouter/free", c["base"], c["key"]):
                                recovered = True
                                full += chunk
                                yield chunk
                        except ModelError:
                            pass
                    if not recovered and not got_any:
                        if e.status in (401, 403):
                            note = f"\n\n⚠️ کلید «{c['name']}» نامعتبر یا منقضی است."
                        elif e.status == 429:
                            note = f"\n\n⚠️ ظرفیت رایگان «{c['name']}» فعلاً پر است. کمی بعد دوباره تلاش کنید."
                        elif e.status == 404:
                            note = f"\n\n⚠️ مدل «{c['name']}» در سرویسش پیدا نشد. مدیر باید نام مدل را اصلاح کند."
                        else:
                            note = f"\n\n⚠️ «{c['name']}» فعلاً در دسترس نیست (کد {e.status})."
                        full += note
                        yield note

            if web_note:
                full += web_note
                yield web_note

        except (httpx.ConnectError, httpx.ReadError, httpx.ConnectTimeout, httpx.RemoteProtocolError):
            err = "⚠️ اتصال به سرویس هوش مصنوعی برقرار نشد. کمی بعد دوباره تلاش کنید."
            full += err
            yield err
        finally:
            d = get_db()
            d.execute(
                "INSERT INTO messages (conversation_id, role, content) VALUES (?, 'assistant', ?)",
                (conv_id, full or "…"),
            )
            d.commit()
            d.close()

    return StreamingResponse(
        gen(),
        media_type="text/plain; charset=utf-8",
        headers={"X-Conversation-Id": str(conv_id)},
    )


# ---------------- خروجی Word / PDF و سازنده مقاله بلند ----------------

FONT_CHOICES = ["Vazirmatn", "B Nazanin", "IRANSans", "B Titr", "Tahoma", "Times New Roman", "Calibri", "Arial"]


@app.get("/download/{fname}")
def download_file(fname: str, request: Request):
    require_user(request)
    if not re.fullmatch(r"[A-Za-z0-9._-]+", fname):
        raise HTTPException(400, "نام فایل نامعتبر")
    import export_utils
    path = os.path.join(export_utils.EXPORT_DIR, fname)
    if not os.path.exists(path):
        raise HTTPException(404, "فایل پیدا نشد یا منقضی شده است")
    return FileResponse(path, filename=fname)


@app.post("/api/export")
async def api_export(request: Request):
    """تبدیل یک پاسخ به فایل Word / PDF با فونت و سایز دلخواه"""
    user = require_user(request)
    body = await request.json()
    content = (body.get("content") or "").strip()
    if not content:
        raise HTTPException(400, "متنی برای خروجی وجود ندارد")
    font = body.get("font") or "Vazirmatn"
    if font not in FONT_CHOICES:
        font = "Vazirmatn"
    try:
        size = max(8, min(int(body.get("size") or 14), 36))
    except (TypeError, ValueError):
        size = 14
    formats = body.get("formats") or ["docx"]
    title = (body.get("title") or "").strip() or None

    import export_utils
    blocks = export_utils.md_to_blocks(content)
    files, notes = [], []
    if "docx" in formats:
        name = await run_in_threadpool(export_utils.build_docx, blocks, font, size, title)
        files.append({"label": "📄 دانلود Word", "url": f"/download/{name}"})
    if "pdf" in formats:
        name, err = await run_in_threadpool(export_utils.build_pdf, blocks, size, title)
        if name:
            files.append({"label": "📕 دانلود PDF", "url": f"/download/{name}"})
        elif err:
            notes.append(err)
    return {"files": files, "notes": notes}


async def _call_model_once(c, prompt: str, system: str | None = None, max_tokens: int = 1800) -> str:
    """یک فراخوانی بدون استریم — با دو تلاش مجدد در صورت شلوغی"""
    messages = ([{"role": "system", "content": system}] if system else []) + [
        {"role": "user", "content": prompt}
    ]
    payload = {"model": c["model"], "messages": messages, "max_tokens": max_tokens}
    headers = {"Authorization": f"Bearer {c['key']}", "Content-Type": "application/json"}
    for attempt in range(3):
        async with httpx.AsyncClient(timeout=httpx.Timeout(180, connect=15)) as client:
            r = await client.post(f"{c['base']}/chat/completions", json=payload, headers=headers)
        if r.status_code == 200:
            data = r.json()
            return ((data.get("choices") or [{}])[0].get("message") or {}).get("content", "") or ""
        if r.status_code == 429 and attempt < 2:
            await asyncio.sleep(20)
            continue
        raise ModelError(r.status_code, r.text[:300])
    return ""


@app.post("/api/longdoc")
async def api_longdoc(request: Request):
    """سازنده مقاله بلند: فهرست بخش‌ها → نوشتن بخش‌به‌بخش → خروجی Word/PDF"""
    user = require_user(request)
    body = await request.json()
    topic = (body.get("topic") or "").strip()
    if not topic:
        raise HTTPException(400, "موضوع مقاله را بنویسید")
    try:
        pages = max(1, min(int(body.get("pages") or 10), 500))
    except (TypeError, ValueError):
        pages = 10
    font = body.get("font") or "Vazirmatn"
    if font not in FONT_CHOICES:
        font = "Vazirmatn"
    try:
        size = max(8, min(int(body.get("size") or 14), 36))
    except (TypeError, ValueError):
        size = 14
    formats = body.get("formats") or ["docx"]

    catalog = get_ai_catalog()
    c = catalog[0]
    if not c.get("key"):
        raise HTTPException(400, "ابتدا در پنل مدیریت، کلید API را تنظیم کنید")

    n_sections = max(3, min(pages // 2 + 1, 250))
    sys_prompt = BASE_SYSTEM_PROMPT

    # ثبت در تاریخچه گفتگوها
    db = get_db()
    cur = db.execute(
        "INSERT INTO conversations (user_id, title) VALUES (?, ?)",
        (user["id"], f"📄 مقاله: {topic[:45]}"),
    )
    conv_id = cur.lastrowid
    db.execute(
        "INSERT INTO messages (conversation_id, role, content) VALUES (?, 'user', ?)",
        (conv_id, f"📄 درخواست مقاله {pages} صفحه‌ای درباره: {topic}"),
    )
    db.commit()
    db.close()

    async def gen():
        full_log = ""

        def log(t):
            nonlocal full_log
            full_log += t
            return t

        try:
            yield log(f"📄 **ساخت مقاله «{topic}» — حدود {pages} صفحه**\n\n⏳ گام ۱: طراحی فهرست بخش‌ها…\n")
            outline = await _call_model_once(
                c,
                f"برای یک مقاله جامع {pages} صفحه‌ای فارسی درباره «{topic}» دقیقاً {n_sections} عنوان بخش بنویس. "
                "هر عنوان در یک خط جداگانه، بدون شماره و بدون توضیح اضافه. عنوان‌ها متنوع و بدون هم‌پوشانی باشند.",
                system=sys_prompt,
                max_tokens=1200,
            )
            titles = [re.sub(r"^[\d\-.،*#)\s]+", "", t).strip() for t in outline.split("\n")]
            titles = [t for t in titles if 2 < len(t) < 120][:n_sections]
            if not titles:
                yield log("\n⚠️ فهرست بخش‌ها ساخته نشد. دوباره تلاش کنید.")
                return

            yield log(f"✅ {len(titles)} بخش طراحی شد.\n\n")
            article = f"# {topic}\n"
            done_titles = []

            for i, t in enumerate(titles, 1):
                if await request.is_disconnected():
                    yield log("\n⏹ ساخت مقاله توسط کاربر متوقف شد.")
                    break
                yield log(f"⏳ نوشتن بخش {i} از {len(titles)}: «{t}»…\n")
                try:
                    part = await _call_model_once(
                        c,
                        f"مقاله‌ای فارسی درباره «{topic}» در حال نگارش است.\n"
                        f"بخش‌های نوشته‌شده تاکنون: {'، '.join(done_titles) if done_titles else 'هیچ'}.\n"
                        f"اکنون فقط بخش «{t}» را بنویس: حدود ۶۰۰ تا ۸۰۰ کلمه، علمی و ساختارمند. "
                        "از تکرار مطالب و واژه‌های بخش‌های قبلی جداً پرهیز کن و مطالب و واژگان کاملاً تازه بیاور. "
                        "خودِ عنوان بخش را ننویس؛ فقط متن.",
                        system=sys_prompt,
                    )
                except ModelError as e:
                    yield log(f"⚠️ بخش «{t}» به دلیل خطای سرویس (کد {e.status}) رد شد.\n")
                    continue
                article += f"\n\n## {t}\n\n{part.strip()}"
                done_titles.append(t)
                await asyncio.sleep(1)

            yield log("\n⏳ گام پایانی: ساخت فایل‌ها…\n")
            import export_utils
            blocks = export_utils.md_to_blocks(article)
            links = []
            if "docx" in formats:
                name = await run_in_threadpool(export_utils.build_docx, blocks, font, size, topic)
                links.append(f"[📄 دانلود Word](/download/{name})")
            if "pdf" in formats:
                name, err = await run_in_threadpool(export_utils.build_pdf, blocks, size, topic)
                if name:
                    links.append(f"[📕 دانلود PDF](/download/{name})")
                elif err:
                    yield log(f"⚠️ {err}\n")
            yield log(f"\n✅ **مقاله آماده شد!** ({len(done_titles)} بخش)\n\n" + "  |  ".join(links))
        except ModelError as e:
            yield log(f"\n⚠️ سرویس هوش مصنوعی خطا داد (کد {e.status}). کمی بعد دوباره تلاش کنید.")
        except Exception:
            yield log("\n⚠️ خطای غیرمنتظره در ساخت مقاله.")
        finally:
            d = get_db()
            d.execute(
                "INSERT INTO messages (conversation_id, role, content) VALUES (?, 'assistant', ?)",
                (conv_id, full_log or "…"),
            )
            d.commit()
            d.close()

    return StreamingResponse(
        gen(),
        media_type="text/plain; charset=utf-8",
        headers={"X-Conversation-Id": str(conv_id)},
    )


# ---------------- API پنل مدیریت ----------------

@app.get("/admin/ai_settings")
def admin_get_ai_settings(request: Request):
    require_admin(request)
    db = get_db()
    row = db.execute("SELECT value FROM settings WHERE key = 'ai_config'").fetchone()
    db.close()
    if row:
        try:
            return json.loads(row["value"])
        except json.JSONDecodeError:
            pass
    return {"provider": PROVIDER, "api_key": "", "model": "", "ais": []}


@app.post("/admin/ai_settings")
async def admin_save_ai_settings(request: Request):
    require_admin(request)
    body = await request.json()
    cfg = {
        "provider": (body.get("provider") or "groq").strip(),
        "api_key": (body.get("api_key") or "").strip(),
        "model": (body.get("model") or "").strip(),
        "ais": [
            {
                "name": (a.get("name") or "").strip(),
                "service": (a.get("service") or "").strip(),
                "model": (a.get("model") or "").strip(),
                "key": (a.get("key") or "").strip(),
            }
            for a in (body.get("ais") or [])[:10]
        ],
    }
    db = get_db()
    db.execute(
        "INSERT INTO settings (key, value) VALUES ('ai_config', ?) "
        "ON CONFLICT(key) DO UPDATE SET value = excluded.value",
        (json.dumps(cfg, ensure_ascii=False),),
    )
    db.commit()
    db.close()
    return {"ok": True, "models": len(get_ai_catalog())}


@app.post("/admin/test_ai")
async def admin_test_ai(request: Request):
    """تست زنده یک کلید — یک پیام کوتاه به سرویس می‌فرستد"""
    require_admin(request)
    body = await request.json()
    base = resolve_base(body.get("service"))
    model = (body.get("model") or "").strip()
    key = (body.get("key") or "").strip()
    if not (base and model and key):
        return {"ok": False, "msg": "سرویس، نام مدل و کلید را کامل وارد کنید"}
    payload = {"model": model, "messages": [{"role": "user", "content": "سلام"}], "max_tokens": 10}
    try:
        async with httpx.AsyncClient(timeout=httpx.Timeout(30, connect=10)) as client:
            r = await client.post(
                f"{base}/chat/completions",
                json=payload,
                headers={"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
            )
        if r.status_code == 200:
            return {"ok": True, "msg": "✅ کلید سالم است و مدل پاسخ داد"}
        if r.status_code in (401, 403):
            return {"ok": False, "msg": "❌ کلید نامعتبر یا منقضی است"}
        if r.status_code == 404:
            return {"ok": False, "msg": "❌ نام مدل در این سرویس پیدا نشد"}
        if r.status_code == 429:
            return {"ok": False, "msg": "⚠️ کلید درست است ولی ظرفیت رایگان فعلاً پر است"}
        return {"ok": False, "msg": f"❌ خطای سرویس (کد {r.status_code})"}
    except Exception:
        return {"ok": False, "msg": "❌ اتصال به سرویس برقرار نشد"}


@app.get("/admin/data")
def admin_data(request: Request):
    require_admin(request)
    db = get_db()
    codes = db.execute("SELECT * FROM codes ORDER BY id DESC LIMIT 500").fetchall()
    users = db.execute(
        "SELECT id, username, stars, is_admin, device_fp, code_used, created_at FROM users ORDER BY id DESC"
    ).fetchall()
    db.close()
    return {"codes": [dict(c) for c in codes], "users": [dict(u) for u in users]}


@app.post("/admin/codes")
async def admin_generate_codes(request: Request):
    require_admin(request)
    body = await request.json()
    stars = int(body.get("stars", 1))
    count = max(1, min(int(body.get("count", 1)), 200))
    if stars not in (1, 2, 3, 4):
        raise HTTPException(400, "ستاره باید بین ۱ تا ۴ باشد")
    db = get_db()
    new_codes = []
    for _ in range(count):
        code = generate_code()
        db.execute("INSERT INTO codes (code, stars) VALUES (?, ?)", (code, stars))
        new_codes.append(code)
    db.commit()
    db.close()
    return {"codes": new_codes, "stars": stars}


@app.post("/admin/create_user")
async def admin_create_user(request: Request):
    """ساخت دستی کاربر توسط ادمین (بدون نیاز به کد)"""
    require_admin(request)
    body = await request.json()
    username = (body.get("username") or "").strip()
    password = body.get("password") or ""
    stars = int(body.get("stars", 1))
    if len(username) < 3 or len(password) < 6 or stars not in (1, 2, 3, 4):
        raise HTTPException(400, "اطلاعات نامعتبر (نام ≥ ۳ حرف، گذرواژه ≥ ۶ حرف، ستاره ۱ تا ۴)")
    db = get_db()
    if db.execute("SELECT 1 FROM users WHERE username = ?", (username,)).fetchone():
        db.close()
        raise HTTPException(400, "این نام کاربری وجود دارد")
    db.execute(
        "INSERT INTO users (username, password, stars) VALUES (?, ?, ?)",
        (username, hash_pw(password), stars),
    )
    db.commit()
    db.close()
    return {"ok": True}


@app.post("/admin/reset_device")
async def admin_reset_device(request: Request):
    """آزاد کردن قفل دستگاه یک کاربر (برای انتقال به گوشی/کامپیوتر جدید)"""
    require_admin(request)
    body = await request.json()
    db = get_db()
    db.execute("UPDATE users SET device_fp = NULL WHERE id = ?", (int(body["user_id"]),))
    db.commit()
    db.close()
    return {"ok": True}


@app.post("/admin/set_stars")
async def admin_set_stars(request: Request):
    require_admin(request)
    body = await request.json()
    stars = int(body.get("stars", 1))
    if stars not in (1, 2, 3, 4):
        raise HTTPException(400, "ستاره باید بین ۱ تا ۴ باشد")
    db = get_db()
    db.execute("UPDATE users SET stars = ? WHERE id = ? AND is_admin = 0", (stars, int(body["user_id"])))
    db.commit()
    db.close()
    return {"ok": True}


@app.post("/admin/delete_user")
async def admin_delete_user(request: Request):
    require_admin(request)
    body = await request.json()
    db = get_db()
    db.execute("DELETE FROM users WHERE id = ? AND is_admin = 0", (int(body["user_id"]),))
    db.commit()
    db.close()
    return {"ok": True}
